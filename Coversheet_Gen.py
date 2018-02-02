'''
Cover Sheet Generator:
Produces a word docx coversheet for homework assignments

github repository: https://github.com/CFreeman217/coversheet.git
'''
import csv
import calendar
import docx
# The built-in comma-separated-value library is useful
# for parsing data files.


REF_WORD_DOC = docx.Document('pyCoverSheet.docx')

class UmkcCourse:
    '''
    Represents an instance of the UMKC Class that I am taking this semester.
    '''
    def __init__(self, prefix, c_num, c_title, instructor, sem, year, count):
        self.prefix = prefix
        self.course_number = c_num
        self.title = c_title
        self.professor = instructor
        self.semester = sem
        self.year = year
        self.count = count
    @property
    def printname(self):
        '''Generates a string of course information in a pretty format '''
        return '{} {} : {}'.format(self.prefix, self.course_number, self.title)

    @property
    def instructorterm(self):
        '''Generates string of instructor and term information for the second line of text'''
        return '{} - {} {}'.format(COURSEDICT[COURSE_SELECT].professor,
                                   COURSEDICT[COURSE_SELECT].semester,
                                   COURSEDICT[COURSE_SELECT].year)
    @property
    def savestring(self):
        ''' Creates a string for saving into a csv format'''
        return '{},{},{},{},{},{},{}\n'.format(self.prefix,
                                               self.course_number,
                                               self.title,
                                               self.professor,
                                               self.semester,
                                               self.year,
                                               self.count)

    def addfile(self, filename):
        ''' Writes the save string in CSV format to the file name passed '''
        with open(filename, 'a') as savefile:
            savefile.write(self.savestring)

def user_courseinput(in_num):
    '''Uses input prompts to generate a UI_NEW_COURSE_SELECT instance of the class UmkcClass'''
    pfx = input('Enter Prefix [ME] : ')
    if pfx == '':
        pfx = 'ME'
    n_course = input('Enter Course Number [{}]: '.format(in_num))
    if n_course == '':
        n_course = str(in_num)
    t_course = input('Enter Course Title : ').title()
    i_course = input('Enter Instructor Name :').title()
    s_course = input('Enter Semester [(f)/s]: ').title()
    if s_course == '':
        s_course = 'Fall'
    y_course = input('Enter Year (2018) : ')
    if y_course == '':
        y_course = '2018'
    c_course = 0
    COURSEDICT[n_course] = UmkcCourse(pfx,
                                      n_course,
                                      t_course,
                                      i_course,
                                      s_course,
                                      y_course,
                                      c_course)
    COURSEDICT[n_course].addfile('savefile.csv')

def display_courseload():
    '''Generates a display of the currently loaded courses in the savefile csv'''
    print('\nCurrently Loaded Courses : \n')
    for dictkey, course in COURSEDICT.items():
        print('{} : {}'.format(dictkey, course.printname))

# Function parses tab delimited file information
# filename must be within the current working directory
def readfile(filename):
    '''Reads CSV file information and in this case, passes it into the course class'''
    c_dict = {}
    # Using a 'with' statement is safer than
    # needing to remember to close the file afer reading
    with open(filename) as file:
        # Delimiter is the character separating the values
        reader = csv.reader(file)
        # Generates a list of the stored information
        data = list(reader)
    # Returns the list of data gathered from the file
    for entry in data:
        c_dict[entry[1]] = UmkcCourse(entry[0],
                                      entry[1],
                                      entry[2],
                                      entry[3],
                                      entry[4],
                                      entry[5],
                                      entry[6])
    return c_dict

def stripdate(in_str):
    '''Cuts the input name into desired date format.
    Uses the calendar built-in library to produce a string month.'''
    month = int(in_str[:2])
    day = int(in_str[2:4])
    year = int('20' + in_str[4:])
    return '{} {} {}'.format(day, calendar.month_name[month], year)


COURSEDICT = readfile('savefile.csv')


display_courseload()
COURSE_SELECT = input('Generate Homework Coversheet for class : ')
if COURSE_SELECT not in COURSEDICT.keys():
    UI_NEW_COURSE_SELECT = input('New Course Number Detected...\n\tCreate New Entry? [(y)/n] : ')
    if UI_NEW_COURSE_SELECT == '' or UI_NEW_COURSE_SELECT.lower() == 'y':
        user_courseinput(COURSE_SELECT)
ASSIGNMENT_NUMBER = input('Assignment Number ({}) : '.format(COURSEDICT[COURSE_SELECT].count))
if ASSIGNMENT_NUMBER == '':
    ASSIGNMENT_NUMBER = COURSEDICT[COURSE_SELECT].count
while True:
    DUEDATE = input('Enter Due Date [MMDDYY] : ')
    if DUEDATE.isdigit() and len(DUEDATE) == 6:
        DUEDATE_STRING = stripdate(DUEDATE)
        break
    print('Enter Date in the form : MMDDYY')








REF_WORD_DOC.add_paragraph(COURSEDICT[COURSE_SELECT].printname, 'Subtitle')
REF_WORD_DOC.add_paragraph(COURSEDICT[COURSE_SELECT].instructorterm, 'Subtitle')
REF_WORD_DOC.add_paragraph('Homework Assignment No. {}'.format(ASSIGNMENT_NUMBER), 'Subtitle')
REF_WORD_DOC.add_paragraph(DUEDATE_STRING, 'Subtitle')
REF_WORD_DOC.save('ME{}_HW{}_coversheet_{}.docx'.format(COURSE_SELECT, ASSIGNMENT_NUMBER, DUEDATE))
