'''
Cover Sheet Generator:
Produces a word docx coversheet for homework assignments

FUTURE WORK:
 - Collect date and schedule information and make a guess for the intended assignment due date.
 - Implement a GUI

github repository: https://github.com/CFreeman217/coversheet.git
'''
import json
import docx

def stripdate(in_str):
    import calendar
    '''Cuts the input name into desired date format.
    Uses the calendar built-in library to produce a string month.'''
    month = int(in_str[:2])
    day = int(in_str[2:4])
    year = int('20' + in_str[4:])
    return '{} {} {}'.format(day, calendar.month_name[month], year)

def savedict(in_dict, filename):
    with open(filename,'w') as sfile:
        sfile.write(json.dumps(in_dict))
    return

def get_title_width(in_dict):
    return max((len(in_dict[k]['c_Title']), k) for k in in_dict)

def course_display(inputfile):
    disptitle = 'UMKC MECHANICAL ENGINEERING'
    maxwidth = len(disptitle)
    leftcol = 10
    with open(inputfile,'r') as rfile:
        coursedata = rfile.read()
        c_dict = json.loads(coursedata)
    datalength = get_title_width(c_dict)
    totallength = leftcol + datalength[0]
    if totallength > maxwidth:
        maxwidth = totallength
    print(disptitle.center(maxwidth, '='))
    for course in c_dict:
        dispnum = '[{}]'.format(course)
        print(dispnum.center(leftcol) + c_dict[course]['c_Title'].rjust(datalength[0]))
    return c_dict

def gather_input(inputfile,templatefile):
    courseinfo = course_display(inputfile)
    while True:
        c_select = input('Print coversheet for which course ? ')
        if c_select.isdigit():
            break
        print('Enter course number as an integer')
    if c_select not in courseinfo.keys():
        addcourse = input('New Course Number Detected...\n\tCreate New Entry? [(y)/n] : ')
        if addcourse == '' or addcourse.lower() == 'y':
            courseinfo[c_select] = {}
            courseinfo[c_select]['c_Title'] = input('Enter Course Name : ').title()
            courseinfo[c_select]['c_num'] = c_select
            courseinfo[c_select]['Instr'] = input('Instructor Name : ').title()
            courseinfo[c_select]['Semes'] = input('Semester [Spring] : ').title()
            if courseinfo[c_select]['Semes'] == '':
                courseinfo[c_select]['Semes'] = 'Spring'
            courseinfo[c_select]['Year'] = input('Year [2018] : ').title()
            if courseinfo[c_select]['Year'] == '':
                courseinfo[c_select]['Year'] = '2018'
            courseinfo[c_select]['Asign'] = ''
            courseinfo[c_select]['Probs'] = ''
            courseinfo[c_select]['Ddate'] = None
    if courseinfo[c_select]['Asign'].isdigit():
        courseinfo[c_select]['Asign'] = str(int(courseinfo[c_select]['Asign']) + 1)
    assign = input('Enter Assignment Name or Number [{}] : '.format(courseinfo[c_select]['Asign']))
    if assign != '':
        courseinfo[c_select]['Asign'] = assign
    if c_select == '385':
        if assign == '6':
            courseinfo[c_select]['Probs'] = 'CH 4 : 5, 10, 15, 19. 25, 32, 42'
            courseinfo[c_select]['Ddate'] = '030618'
        elif assign == '7':
            courseinfo[c_select]['Probs'] = 'CH 4 : 52d. 56, 65. 91       CH 5 : 2. 4, 16. 18'
            courseinfo[c_select]['Ddate'] = '031318'
        elif assign == '8':
            courseinfo[c_select]['Probs'] = 'CH 5 : 28. 36, 44, 56'
            courseinfo[c_select]['Ddate'] = '032018'
        elif assign == '9':
            courseinfo[c_select]['Probs'] = 'CH 6 : 6, 12, 32, 39, 48, 56'
            courseinfo[c_select]['Ddate'] = '041218'
        elif assign == '10':
            courseinfo[c_select]['Probs'] = 'CH 7 : 3, 8, 15, 28, 35, 48, 52, 57'
            courseinfo[c_select]['Ddate'] = '050118'
    else:
        oldprobs = courseinfo[c_select]['Probs']
        courseinfo[c_select]['Probs'] = input('Enter Problem Info [{}] : '.format(courseinfo[c_select]['Probs']))
        if oldprobs != '' and courseinfo[c_select]['Probs'] == '':
            use_old = input('Use existing problem data? [n]')
            if use_old.lower == 'y':
                courseinfo[c_select]['Probs'] = oldprobs
        while True:
            duedate = input('Enter Due Date [MMDDYY] : ')
            if duedate.isdigit() and len(duedate) == 6:
                courseinfo[c_select]['Ddate'] = duedate
                break
            print('Enter Date in the form : MMDDYY')
    savedict(courseinfo, inputfile)
    create_coversheet(courseinfo[c_select], templatefile)

def create_coversheet(course_in, t_file):
    import docx
    cs_doc = docx.Document(t_file)
    c_num = course_in['c_num']
    cs_doc.add_paragraph('Clay Freeman', 'Title')
    line_1 = 'ME {} : {}'.format(c_num, course_in['c_Title'])
    cs_doc.add_paragraph(line_1, 'Subtitle')
    line_2 = '{} - {} {}'.format(course_in['Instr'], course_in['Semes'], course_in['Year'])
    cs_doc.add_paragraph(line_2, 'Subtitle')
    if course_in['Asign'].isdigit():
        line_3 = 'Homework Assignment {}'.format(course_in['Asign'])
    else:
        line_3 = '{} Homework'.format(course_in['Asign'])
    cs_doc.add_paragraph(line_3, 'Subtitle')
    if course_in['Probs'] != '':
        newline = '{}'.format(course_in['Probs'])
        cs_doc.add_paragraph(newline, 'Subtitle')
    cs_doc.add_paragraph(stripdate(course_in['Ddate']), 'Subtitle')
    cs_doc.save('ME{}_HW{}_coversheet_{}.docx'.format(c_num, course_in['Asign'],course_in['Ddate']))

gather_input('.courseinfo.txt','pyCoverSheet.docx')











