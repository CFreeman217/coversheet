#!/usr/bin/env python
# https://github.com/CFreeman217/coversheet.git
import docx
import calendar
doc = docx.Document('pyCoverSheet.docx')

courses = {'380' : ('ME 380 : Manufacturing Methods', 'Prof. B. Hanlin'),
            '352' : ('ME 352 : Instrumentation and Measurements', 'Prof. J. Mahoney'),
            '306' : ('ME 306 : Computer Aided Engineering', 'Dr. A. Stylianou'),
            '385' : ('ME 385 : System Dynamics', 'Dr. D. Justice'),
            '399' : ('ME 399 : Heat and Mass Transfer', 'Prof. J. Mahoney')}

def stripDate(in_str):
    month = int(in_str[:2])
    day = int(in_str[2:4])
    year = int('20' + in_str[4:])
    return '{} {} {}'.format(day, calendar.month_name[month], year)
    

print('\n Current Courses: \n')
for ckey in courses.keys():
    print(ckey + ' : ' + courses[ckey][0] )
c_select = input('Generate Homework Coversheet for class : ')
a_num = input('Assignment Number : ')
# problems = input('Assignment Problem Numbers (String) (Optional) : \n')
date_in = input('Due date MMDDYY : ')
duedate = stripDate(date_in)


line_1 = courses[c_select][0]
line_2 = '{} - Spring 2018'.format(courses[c_select][1])
line_3 = 'Homework Assignment No. {}'.format(a_num)
# line_4 = problems
line_5 = duedate

doc.add_paragraph(line_1, 'Subtitle')
doc.add_paragraph(line_2, 'Subtitle')
doc.add_paragraph(line_3, 'Subtitle')
doc.add_paragraph(line_5, 'Subtitle')
# doc.paragraphs[1] = line_1
# doc.paragraphs[1].style = 'Subtitle'
# doc.paragraphs[2] = line_2
# doc.paragraphs[2].style = 'Subtitle'
# doc.paragraphs[3] = line_3
# doc.paragraphs[3].style = 'Subtitle'
# doc.paragraphs[4] = line_4
# doc.paragraphs[4].style = 'Subtitle'
# doc.paragraphs[5] = line_5
# doc.paragraphs[5].style = 'Subtitle'

doc.save('ME{}_HW{}_coversheet_{}.docx'.format(c_select, a_num, date_in))